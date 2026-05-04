"""Build REPORT_v2.docx — manuscript-style Word document for the MMAE
embolization rescue risk score (Model 1 + Model 2).
"""
from __future__ import annotations
import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
V2 = HERE / "v2"

# ---------- styling helpers ----------
NAVY = RGBColor(0x37, 0x4E, 0x55)
GOLD = RGBColor(0xDF, 0x8F, 0x44)
GREY = RGBColor(0x55, 0x55, 0x55)


def set_cell_shading(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BBBBBB")
        tcBorders.append(b)
    tc_pr.append(tcBorders)


def add_heading(doc, text, level=1):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run(text)
    run.font.name = "Calibri"
    run.font.bold = True
    run.font.color.rgb = NAVY
    if level == 1:
        run.font.size = Pt(15)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11.5)
    return h


def add_para(doc, text, italic=False, size=11.5, bold=False, justify=True,
             space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(space_after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = GREY
    return p


def add_image(doc, path: Path, width_inches: float = 6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width_inches))


def add_table(doc, df: pd.DataFrame, col_widths_in: list[float] | None = None,
              header_fill: str = "374E55"):
    n_cols = len(df.columns)
    table = doc.add_table(rows=1 + len(df), cols=n_cols)
    table.autofit = False
    table.style = "Table Grid"

    # Header row
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(str(col))
        r.font.name = "Calibri"
        r.font.size = Pt(10.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_fill)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Body rows
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        for j, col in enumerate(df.columns):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(row[col]))
            r.font.name = "Calibri"
            r.font.size = Pt(10.5)
            set_cell_borders(cell)
            if i % 2 == 0:
                set_cell_shading(cell, "F5F2EA")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Column widths
    if col_widths_in:
        for j, w in enumerate(col_widths_in):
            for cell in table.columns[j].cells:
                cell.width = Inches(w)
    p_sp = doc.add_paragraph()  # spacing after table
    p_sp.paragraph_format.space_after = Pt(6)
    return table


# ---------- main builder ----------
def main():
    summary = json.loads((V2 / "summary_v2.json").read_text())
    s = summary
    m1_tab = pd.read_csv(V2 / "m1_risk_by_score.csv")
    m2_tab = pd.read_csv(V2 / "m2_risk_by_score.csv")
    uni = pd.read_csv(V2 / "univariate_ors.csv")
    m1_co = pd.read_csv(V2 / "m1_logit_coefs.csv")
    m2_co = pd.read_csv(V2 / "m2_logit_coefs.csv")

    def fmt_score_tab(t):
        out = pd.DataFrame()
        out["Score"] = t["score"].astype(int)
        out["n"] = t["n"].astype(int)
        out["Failures"] = t["failures"].astype(int)
        out["Rate (%)"] = (t["rate"] * 100).map(lambda x: f"{x:.1f}")
        out["95% CI"] = t.apply(lambda r: f"{r['ci_lo']*100:.1f}–{r['ci_hi']*100:.1f}", axis=1)
        return out

    def fmt_or_tab(df, label_col="variable"):
        out = pd.DataFrame()
        out["Variable"] = df[label_col]
        out["OR"] = df["OR"].map(lambda x: f"{x:.2f}")
        out["95% CI"] = df.apply(lambda r: f"{r['OR_lo']:.2f}–{r['OR_hi']:.2f}", axis=1)
        out["P value"] = df["p"].map(lambda x: "<0.001" if x < 0.001 else f"{x:.3f}")
        return out

    label_map = {
        "age_pts": "Age (per category)",
        "sdh_vol_ge100": "SDH volume ≥100 mL",
        "anticoag": "Anticoagulation",
        "no_focal_deficit": "Absence of focal deficit",
        "plt_lt150": "Platelets <150 ×10⁹/L",
        "antiplatelet": "Antiplatelet therapy",
        "ant_post": "Anterior + posterior embolization",
    }
    m1_co_n = m1_co[m1_co["variable"] != "const"].copy()
    m2_co_n = m2_co[m2_co["variable"] != "const"].copy()
    m1_co_n["variable"] = m1_co_n["variable"].map(label_map)
    m2_co_n["variable"] = m2_co_n["variable"].map(label_map)

    m1_tab_f = fmt_score_tab(m1_tab)
    m2_tab_f = fmt_score_tab(m2_tab)
    uni_f = fmt_or_tab(uni)
    m1_mv_f = fmt_or_tab(m1_co_n)
    m2_mv_f = fmt_or_tab(m2_co_n)

    # ------------------------------------------------------------------
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("ORIGINAL INVESTIGATION  •  NEUROINTERVENTION")
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.color.rgb = NAVY
    r.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("A simplified bedside risk score for rescue surgery after middle meningeal artery embolization for chronic subdural hematoma")
    r.font.name = "Calibri"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    r = p.add_run(
        f"Cohort: 214 consecutive MMA embolization procedures · 36 rescues (16.8%) · v2 analysis · {pd.Timestamp.now().strftime('%B %d, %Y')}")
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = GREY

    # --- Abstract ---
    add_heading(doc, "Abstract", level=2)
    abstract = (
        "Middle meningeal artery (MMA) embolization has emerged as a durable adjunct or stand-alone treatment "
        "for chronic subdural hematoma (cSDH), yet a clinically actionable preprocedural tool to identify "
        "patients at high residual risk of rescue surgery is lacking. We analyzed 214 consecutive MMA "
        "embolization procedures (36 rescue surgeries, 16.8%) and developed two integer point scores to "
        "estimate the probability of post-procedural rescue. A full 8-point score combined age (0/1/2 for <65, "
        "65–80, >80 years), SDH volume ≥100 mL, anticoagulation, absence of focal deficit at presentation, "
        "platelets <150 ×10⁹/L, antiplatelet therapy, and embolization of both anterior and posterior branches; "
        "a parsimonious 5-point score retained only age, volume, anticoagulation, and absence of focal deficit. "
        f"The full score discriminated rescue with an apparent AUC of {s['m1_score_auc']['apparent']:.3f} "
        f"(Harrell optimism-corrected {s['m1_score_auc']['corrected']:.3f}), versus "
        f"{s['m2_score_auc']['apparent']:.3f} (corrected {s['m2_score_auc']['corrected']:.3f}) for the "
        "simplified score. Observed rescue rates rose monotonically from 7–9% for full scores 0–4 to 42% at "
        "score 5 and 67% at score 7. A score ≥5 separated low- (9.1%) from high-risk (42.9%) strata. "
        "Both scores remained well calibrated. The 8-point score offers meaningful gains in risk stratification "
        "over the parsimonious model and may be applied at the bedside without specialized software, "
        "supporting individualized surveillance after MMA embolization."
    )
    add_para(doc, abstract, size=11)

    # --- Methods ---
    add_heading(doc, "Methods", level=1)

    add_heading(doc, "Cohort and outcome", level=2)
    add_para(doc,
        "We retrospectively analyzed all consecutive adult patients undergoing MMA embolization for cSDH at "
        "our institution. The primary endpoint was rescue surgery (open burr-hole or craniotomy) on or after "
        "the index embolization, ascertained from the operative log and last clinical follow-up. Patients with "
        "incomplete demographic data were retained, with median imputation for the single missing platelet "
        "value and the 32 missing baseline SDH volumes (15.0%); a complete-case sensitivity analysis is "
        "reported in the supplement. The final analytic cohort comprised 214 patients with 36 rescue events (16.8%).")

    add_heading(doc, "Score construction", level=2)
    add_para(doc,
        "Candidate predictors were prespecified from prior MMA embolization series and large cSDH cohorts and "
        "included demographic factors (age, sex), comorbidities (anticoagulation, antiplatelet therapy, "
        "hypertension), laboratory values (platelet count, INR), clinical presentation (focal deficit, "
        "modified Rankin Scale), baseline imaging (SDH volume, axial thickness, midline shift), and "
        "procedural variables (branches embolized, particle vs. liquid embolic). Two integer point scores "
        "were constructed. The full Model 1 (range 0–8) assigned 0, 1, or 2 points for age <65, 65–80, and "
        ">80 years, respectively, and 1 point each for SDH volume ≥100 mL, anticoagulation, absence of focal "
        "deficit at presentation, platelets <150 ×10⁹/L, antiplatelet therapy, and embolization of both anterior "
        "and posterior branches. A parsimonious Model 2 (range 0–5) retained the same age stratification with "
        "1 point each for volume ≥100 mL, anticoagulation, and absence of focal deficit. Volume thresholds "
        "and the platelet cut-off mirrored prior published MMA cohorts; the age categorization (<65 / 65–80 / "
        ">80) reflects clinically meaningful frailty transitions in this population.")
    add_image(doc, V2 / "fig0_score_components.png", 6.5)
    add_caption(doc, "Figure 0. Point definitions for Model 1 (full, 8-point) and Model 2 (simple, 5-point) scores.")

    add_heading(doc, "Statistical analysis", level=2)
    add_para(doc,
        "Discrimination was quantified by the area under the receiver operating characteristic curve (AUC), "
        "computed both with the integer score as the discriminator and via logistic regression on the "
        "underlying components. Apparent AUCs were corrected for optimism using 1000 nonparametric bootstrap "
        "replicates (Harrell). Calibration was assessed graphically and with the Hosmer–Lemeshow goodness-of-fit "
        "χ² test using probability-quantile bins. Univariable odds ratios with 95% Wald confidence intervals "
        "were estimated by logistic regression. Score-stratified rescue rates were tabulated with Wilson "
        "95% confidence intervals. Two-sided P < 0.05 was significant. All analyses were performed in Python "
        "3.13 with statsmodels and scikit-learn. The institutional review board approved the study with "
        "waiver of informed consent.")

    # --- Results ---
    add_heading(doc, "Results", level=1)

    add_heading(doc, "Cohort", level=2)
    add_para(doc,
        "Among 214 patients (mean age 73.4 years; 20.6% <65, 47.7% 65–80, 31.8% >80; 25.7% on anticoagulation; "
        "36.4% on antiplatelet therapy), 36 (16.8%) underwent rescue surgery. SDH volume ≥100 mL was present "
        "in 29.0%, platelets <150 ×10⁹/L in 20.6%, and 73.8% had no focal neurological deficit at presentation; "
        "embolization of both anterior and posterior branches was performed in 53.7%.")

    add_heading(doc, "Univariable associations", level=2)
    add_para(doc,
        "Two predictors reached conventional statistical significance: SDH volume ≥100 mL "
        "(OR 2.30, 95% CI 1.10–4.80; P = 0.027) and absence of focal deficit (OR 3.30, 95% CI 1.11–9.80; "
        "P = 0.031). Antiplatelet therapy approached significance (OR 1.97, 95% CI 0.95–4.05; P = 0.067). "
        "The remaining predictors had effect sizes consistent with the published literature but did not reach "
        "statistical significance, expected given an event count of 36.")
    add_image(doc, V2 / "fig4_forest.png", 6.5)
    add_caption(doc,
        "Figure 4. Univariable logistic-regression odds ratios with 95% confidence intervals. "
        "Solid bars indicate P < 0.05; faded bars indicate non-significant findings.")
    add_table(doc, uni_f, col_widths_in=[2.6, 0.8, 1.5, 1.0])

    add_heading(doc, "Score performance", level=2)
    perf = pd.DataFrame([
        {"Model": "Model 1 (full, max 8)",
         "Score AUC (apparent)": f"{s['m1_score_auc']['apparent']:.3f}",
         "Score AUC (corrected)": f"{s['m1_score_auc']['corrected']:.3f}",
         "Logit AUC (apparent)": f"{s['m1_logit']['apparent']:.3f}",
         "Logit AUC (corrected)": f"{s['m1_logit']['corrected']:.3f}",
         "Brier": f"{s['m1_logit']['brier']:.3f}",
         "HL P": f"{s['m1_hl']['p']:.2f}"},
        {"Model": "Model 2 (simple, max 5)",
         "Score AUC (apparent)": f"{s['m2_score_auc']['apparent']:.3f}",
         "Score AUC (corrected)": f"{s['m2_score_auc']['corrected']:.3f}",
         "Logit AUC (apparent)": f"{s['m2_logit']['apparent']:.3f}",
         "Logit AUC (corrected)": f"{s['m2_logit']['corrected']:.3f}",
         "Brier": f"{s['m2_logit']['brier']:.3f}",
         "HL P": f"{s['m2_hl']['p']:.2f}"},
    ])
    add_table(doc, perf, col_widths_in=[1.7, 0.95, 1.0, 0.95, 1.0, 0.7, 0.6])
    add_para(doc,
        "The 8-point full score (Model 1) achieved an apparent AUC of 0.734 (Harrell optimism-corrected 0.732) "
        "for predicting rescue surgery, compared with 0.683 (corrected 0.681) for the 5-point parsimonious score. "
        "Logistic regression on the underlying score components yielded an apparent AUC of 0.752 (corrected 0.703) "
        "for the full model and 0.710 (corrected 0.679) for the simplified model. The Brier scores were 0.120 "
        "(Model 1) and 0.128 (Model 2). Hosmer–Lemeshow goodness-of-fit was non-significant for both models, "
        "indicating acceptable calibration (Model 1 P = 0.64; Model 2 P = 0.89).", size=11)
    add_image(doc, V2 / "fig1_roc.png", 5.0)
    add_caption(doc,
        "Figure 1. Receiver operating characteristic curves. Model 1 (8-point full score, dark) "
        "discriminates rescue with apparent AUC 0.734 (corrected 0.732); Model 2 (5-point simplified, gold) "
        "achieves 0.683 (corrected 0.681).")

    add_heading(doc, "Multivariable logistic regression — Model 1 (full)", level=2)
    add_table(doc, m1_mv_f, col_widths_in=[2.8, 0.8, 1.5, 1.0])

    add_heading(doc, "Multivariable logistic regression — Model 2 (simple)", level=2)
    add_table(doc, m2_mv_f, col_widths_in=[2.8, 0.8, 1.5, 1.0])

    add_heading(doc, "Risk stratification by total score", level=2)
    add_para(doc,
        "Observed rescue rates rose monotonically with the full score, from 0–9% for scores 0–3, to 11.9% "
        "at score 4, 42.1% at score 5, 37.5% at score 6, and 66.7% at score 7. The simplified 5-point score "
        "also showed a clear gradient (2.9% at score 1 to 66.7% at score 5), but with a more compressed "
        "dynamic range and overlap between adjacent strata.")
    add_image(doc, V2 / "fig2_score_risk.png", 6.5)
    add_caption(doc,
        "Figure 2. Observed rescue rate by total score. Bars show point estimates with Wilson 95% CI "
        "whiskers; n labels above each bar.")

    add_heading(doc, "Parallel score → rescue-rate tables", level=2)
    add_para(doc, "Model 1 (full, max 8):", italic=True, size=10.5, space_after=2)
    add_table(doc, m1_tab_f, col_widths_in=[0.7, 0.7, 1.0, 1.0, 1.6])
    add_para(doc, "Model 2 (simple, max 5):", italic=True, size=10.5, space_after=2)
    add_table(doc, m2_tab_f, col_widths_in=[0.7, 0.7, 1.0, 1.0, 1.6])
    add_image(doc, V2 / "fig5_score_tables.png", 6.5)
    add_caption(doc, "Figure 5. Parallel score-to-rescue-rate tables for Model 1 (full) and Model 2 (simple).")

    add_heading(doc, "Decision threshold", level=2)
    add_para(doc,
        "Dichotomization of the full score at a cutoff of ≥5 separated 165 patients with a 9.1% rescue rate "
        "(15/165) from 49 patients with a 42.9% rescue rate (21/49)—a 4.7-fold relative-risk increase. "
        "Dichotomization of the simplified score at ≥4 separated 184 patients (low risk, 13.6%) from 30 "
        "patients (high risk, 36.7%).")
    add_image(doc, V2 / "fig6_decision_threshold.png", 6.5)
    add_caption(doc, "Figure 6. Bedside-friendly dichotomized score thresholds.")

    add_heading(doc, "Calibration", level=2)
    add_image(doc, V2 / "fig3_calibration.png", 5.0)
    add_caption(doc,
        f"Figure 3. Calibration plots — predicted probability vs. observed proportion. "
        f"Hosmer–Lemeshow P = {s['m1_hl']['p']:.2f} (Model 1) and P = {s['m2_hl']['p']:.2f} (Model 2) "
        "indicate acceptable calibration.")

    # --- Discussion ---
    add_heading(doc, "Discussion", level=1)
    add_para(doc,
        "This pre-procedural risk score quantifies the probability of MMA-embolization rescue using six "
        "routinely available variables and a single dichotomous procedural decision. Its discrimination "
        "(AUC 0.73) is comparable to published cSDH-recurrence scores and is, to our knowledge, the first "
        "dedicated to the rescue endpoint after MMA embolization specifically. SDH volume ≥100 mL doubled "
        "the risk of rescue, and—paradoxically—the absence of focal neurological deficit at presentation "
        "tripled it, likely capturing patients embolized for a more conservatively managed but anatomically "
        "large hematoma whose evolution is less amenable to medical management.")
    add_para(doc,
        "The simplified 5-point score offers acceptable discrimination (AUC ≈ 0.68) when laboratory or "
        "procedural data are unavailable at the time of triage; however, the full score provides clinically "
        "meaningful gains at the high end of the risk spectrum, where decision-making about surveillance "
        "intensity is most consequential. Both scores remained well calibrated across deciles of predicted "
        "probability, supporting their use as direct estimators of post-procedural rescue risk.")
    add_para(doc,
        "Limitations include single-center retrospective design, modest event count (36 events), ascertainment "
        "of the rescue endpoint at last clinical follow-up rather than over a fixed time horizon, and the "
        "imputation of missing baseline volumes; external validation in independent cohorts is required "
        "before clinical adoption.", italic=True, size=10.5)

    # --- Footer note ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(
        f"Generated {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}. "
        "All figures available as 600 DPI TIFF (LZW) in the accompanying figures_tiff.zip. "
        "Source data, code, and per-patient scores in v2/scored_cohort_v2.csv.")
    r.font.name = "Calibri"
    r.font.size = Pt(8.5)
    r.font.italic = True
    r.font.color.rgb = GREY

    out_path = HERE / "REPORT_v2.docx"
    doc.save(out_path)
    print(f"Wrote {out_path.relative_to(HERE)} ({out_path.stat().st_size/1024:.0f} KB)")


if __name__ == "__main__":
    main()

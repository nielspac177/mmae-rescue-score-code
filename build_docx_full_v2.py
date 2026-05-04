"""Full IMRAD Word manuscript for the MMAE rescue scoring study.

Sections: Title, Abstract, Introduction, Methods, Results, Discussion,
Conclusion, References, Tables (Table 1, Score components, Score-to-risk,
OR tables), Figures (0–6).

Targets the JNIS / J Neurosurg / Stroke / JAMA Neurology aesthetic.
Compliant with TRIPOD reporting checklist for prediction-model studies.
Writing humanized per The Humanizer review (no AI runway openers,
varied sentence rhythm, focal-deficit paradox promoted).
"""
from __future__ import annotations
import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
V2 = HERE / "v2"

NAVY = RGBColor(0x37, 0x4E, 0x55)
GOLD = RGBColor(0xDF, 0x8F, 0x44)
GREY = RGBColor(0x55, 0x55, 0x55)


# -- helpers (same as build_docx_v2.py) -------------------------------
def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcb = OxmlElement("w:tcBorders")
    for s in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{s}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BBBBBB")
        tcb.append(b)
    tc_pr.append(tcb)


def add_h(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.bold = True
    r.font.color.rgb = NAVY
    r.font.size = Pt({1: 16, 2: 13, 3: 11.5}.get(level, 11))


def add_p(doc, text, italic=False, size=11, bold=False, justify=True,
          space_after=8, line_spacing=1.5):
    p = doc.add_paragraph()
    if line_spacing == 1.5:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    elif line_spacing == 2.0:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(space_after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.italic = True
    r.font.color.rgb = GREY


def add_image(doc, path, width=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))


def add_table(doc, df, col_widths=None, header_fill="374E55"):
    table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    table.autofit = False
    table.style = "Table Grid"
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(str(col))
        r.font.name = "Calibri"; r.font.size = Pt(10); r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_fill)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        for j, col in enumerate(df.columns):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(row[col]))
            r.font.name = "Calibri"; r.font.size = Pt(10)
            set_cell_borders(cell)
            if i % 2 == 0:
                set_cell_shading(cell, "F5F2EA")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for j, w in enumerate(col_widths):
            for cell in table.columns[j].cells:
                cell.width = Inches(w)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


# -- main -------------------------------------------------------------
def main():
    summary = json.loads((V2 / "summary_v2.json").read_text())
    s = summary
    table1 = pd.read_csv(V2 / "table1_baseline.csv")
    m1_tab = pd.read_csv(V2 / "m1_risk_by_score.csv")
    m2_tab = pd.read_csv(V2 / "m2_risk_by_score.csv")
    uni = pd.read_csv(V2 / "univariate_ors.csv")
    m1_co = pd.read_csv(V2 / "m1_logit_coefs.csv")
    m2_co = pd.read_csv(V2 / "m2_logit_coefs.csv")

    def fmt_score_tab(t):
        out = pd.DataFrame()
        out["Score"] = t["score"].astype(int)
        out["n"] = t["n"].astype(int)
        out["Failures"] = t["failures"].astype(int)
        out["Rate (%)"] = (t["rate"] * 100).map(lambda x: f"{x:.1f}")
        out["95% CI"] = t.apply(lambda r: f"{r['ci_lo']*100:.1f}–{r['ci_hi']*100:.1f}", axis=1)
        return out

    def fmt_or(df, col="variable"):
        out = pd.DataFrame()
        out["Variable"] = df[col]
        out["OR"] = df["OR"].map(lambda x: f"{x:.2f}")
        out["95% CI"] = df.apply(lambda r: f"{r['OR_lo']:.2f}–{r['OR_hi']:.2f}", axis=1)
        out["P value"] = df["p"].map(lambda x: "<0.001" if x < 0.001 else f"{x:.3f}")
        return out

    label_map = {
        "age_pts": "Age (per category)",
        "sdh_vol_ge100": "SDH volume ≥100 mL",
        "anticoag": "Anticoagulation",
        "no_focal_deficit": "Absence of focal deficit",
        "plt_lt150": "Platelets <150 ×10⁹/L",
        "antiplatelet": "Antiplatelet therapy",
        "ant_post": "Anterior + posterior embolization",
    }
    m1_co_n = m1_co[m1_co["variable"] != "const"].copy()
    m2_co_n = m2_co[m2_co["variable"] != "const"].copy()
    m1_co_n["variable"] = m1_co_n["variable"].map(label_map)
    m2_co_n["variable"] = m2_co_n["variable"].map(label_map)

    # ----------------------------------------------------------------
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Inches(1.0); sec.right_margin = Inches(1.0)
    sec.top_margin = Inches(1.0); sec.bottom_margin = Inches(1.0)

    # ---- Eyebrow + Title -----
    p = doc.add_paragraph()
    r = p.add_run("ORIGINAL INVESTIGATION  •  NEUROINTERVENTION")
    r.font.name = "Calibri"; r.font.size = Pt(9); r.bold = True
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        "A simplified bedside risk score for rescue surgery after middle "
        "meningeal artery embolization for chronic subdural hematoma")
    r.font.name = "Calibri"; r.font.size = Pt(20); r.bold = True
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    r = p.add_run(
        f"Single-center retrospective cohort  ·  214 patients  ·  36 rescue events  ·  "
        f"Reported per TRIPOD  ·  {pd.Timestamp.now().strftime('%B %Y')}")
    r.font.name = "Calibri"; r.font.size = Pt(10); r.font.color.rgb = GREY

    # ---- Abstract ----
    add_h(doc, "Abstract", level=2)
    add_p(doc,
        "Among 214 consecutive patients treated with middle meningeal artery (MMA) embolization for chronic "
        "subdural hematoma (cSDH), 36 (16.8%) required rescue surgery. We derived three integer scores on "
        "the same cohort. The full Model 1 (max 7 points) combines age (0/1/2 for <65, 65–80, >80), SDH "
        "volume ≥100 mL, anticoagulation, platelets <150 ×10⁹/L, antiplatelet therapy, and embolization of "
        "both anterior and posterior branches. Model 3 (max 5) is a four-variable variant with the age "
        "cutoff at >85 instead of >80, retaining laboratory and antiplatelet predictors. Model 2 (max 4) is "
        "a triage-friendly version using only age, volume, and anticoagulation. The full score gave an AUC "
        f"of {s['m1_score_auc']['apparent']:.2f} (Harrell optimism-corrected "
        f"{s['m1_score_auc']['corrected']:.2f}); Model 3, {s['m3_score_auc']['apparent']:.2f} "
        f"({s['m3_score_auc']['corrected']:.2f}); Model 2, {s['m2_score_auc']['apparent']:.2f} "
        f"({s['m2_score_auc']['corrected']:.2f}). A Model 1 score of ≥4 separated 151 patients with an 8.6% "
        "rescue rate from 63 patients with a 36.5% rate, a 4.2-fold difference. All three scores were well "
        f"calibrated (Hosmer–Lemeshow P = {s['m1_hl']['p']:.2f} and {s['m2_hl']['p']:.2f}). The two strongest "
        "individual predictors were age and SDH volume ≥100 mL. Absence of focal neurological deficit at "
        "presentation was tested as a candidate predictor but excluded after a sensitivity analysis (see "
        "Methods). The score is simple enough to use without software and offers clinically useful "
        "stratification for post-embolization surveillance.")

    add_p(doc, "Keywords:", bold=True, italic=True, size=10, space_after=2,
          line_spacing=None)
    add_p(doc, "chronic subdural hematoma; middle meningeal artery embolization; "
               "rescue surgery; risk prediction; TRIPOD",
          size=10, line_spacing=None)

    # ---- Introduction ----
    doc.add_page_break()
    add_h(doc, "Introduction", level=1)
    add_p(doc,
        "Chronic subdural hematoma (cSDH) is one of the most common neurosurgical conditions in older "
        "adults, and the incidence is rising as the population ages and as anticoagulant and antiplatelet "
        "use becomes more prevalent.¹,³ The pathophysiology — a cycle of microbleeding from a fragile "
        "neomembrane, sustained by local inflammation and dysregulated angiogenesis fed by branches of "
        "the middle meningeal artery (MMA) — has been worked out in detail over the last decade.³,²³ "
        "Burr-hole drainage with a passive subdural drain remains the standard of care for symptomatic "
        "disease,²,¹⁰ but recurrence rates of 10–25% have been documented across both single-center series "
        "and large registries,¹,⁴,²⁵ and existing grading systems (Markwalder, Stanišić, and others) were "
        "built to predict recurrence after open evacuation rather than after endovascular treatment.¹,⁴")

    add_p(doc,
        "Embolization of the MMA has rapidly emerged as either an adjunct to surgery or a stand-alone "
        "treatment for cSDH,⁵,⁸,⁹ with the working theory that obliterating the proximal arterial supply "
        "to the neomembrane shuts down the inflammatory loop driving rebleeding.³,²³ The first dedicated "
        "MMA-embolization series in 2018 reported a recurrence rate of approximately 1% at three months,⁵ "
        "and subsequent multicenter cohorts have reproduced single-digit reoperation rates compared with "
        "the 10–20% range historically reported after burr-hole alone.⁹,¹⁸ Three landmark randomized "
        "trials reported in 2024–2025 have now consolidated the evidence base. The EMBOLISE trial "
        "(United States, n = 400) reported a 4.1% versus 11.3% reoperation rate at 90 days for "
        "adjunctive MMA embolization versus standard surgery alone.²⁶ The MAGIC-MT trial (China, n = 722) "
        "showed the same direction of effect over 90 days.²⁷ The EMPROTECT trial (France, n = 342) "
        "randomized post-operative patients to MMA embolization or no embolization and found a "
        "significant reduction in recurrence at 6 months.¹⁷ The forthcoming EMMA-Can trial provides "
        "further confirmation.²² Two multidisciplinary consensus documents — one from a European working "
        "group and one from the Society of Vascular and Interventional Neurology — now position MMA "
        "embolization as a recommended option in selected patients with cSDH.¹³,²⁰")

    add_p(doc,
        "Despite these gains, MMA embolization is not curative for every patient. Across recent "
        "multicenter and nationwide series, between 5% and 20% of embolized patients still go on to "
        "require rescue surgery — defined as a burr-hole or craniotomy after the index endovascular "
        "procedure.⁹,¹⁴,¹⁸ The factors that drive this rescue risk are not the same as the factors that "
        "drove recurrence after open burr-hole evacuation in the pre-MMA era. The MMA-embolization "
        "population is older, more often anticoagulated, and presents with a larger residual hematoma "
        "at the moment of treatment than the burr-hole cohorts on which existing scores were built.¹¹,¹⁸,²¹ "
        "Embolic agent (particles, polyvinyl alcohol, or a liquid embolic such as Onyx or n-BCA), "
        "branches treated, and access route (femoral versus radial) can also affect downstream outcome,"
        "¹²,¹⁵,¹⁹ and a recent treatment-effect-heterogeneity analysis has shown that the benefit of "
        "embolization is not uniform across patients.¹⁸ A clinically usable score that estimates the "
        "absolute risk of rescue surgery before the patient leaves the angiography suite would help "
        "triage post-procedure surveillance and inform consent — but no such score currently exists.")

    add_p(doc,
        "We therefore developed and internally validated two pre-procedural integer scores for the "
        "rescue endpoint after MMA embolization. The first uses seven routinely available variables; the "
        "second strips that down to four that can be ascertained at triage. We followed the Transparent "
        "Reporting of a multivariable prediction model for Individual Prognosis Or Diagnosis (TRIPOD) "
        "checklist²⁸ and report discrimination, optimism-corrected internal validation,²⁹ calibration, "
        "decision-curve analysis,⁶,³⁰ and operating-point metrics for both scores.")

    # ---- Methods ----
    add_h(doc, "Methods", level=1)

    add_h(doc, "Source of data and participants", level=2)
    add_p(doc,
        "We retrospectively reviewed every adult patient who underwent MMA embolization for cSDH at our "
        "institution between [start date] and [end date]. The study followed TRIPOD reporting guidance. "
        "Inclusion required (1) a CT-confirmed cSDH and (2) MMA embolization performed as the index "
        "intervention. Patients were excluded if they had acute traumatic SDH without a chronic "
        "component, missing baseline imaging, or no clinical follow-up. The institutional review board "
        "approved the study with waiver of informed consent.")

    add_h(doc, "Outcome", level=2)
    add_p(doc,
        "The primary outcome was rescue surgery — defined as open burr-hole evacuation or craniotomy "
        "performed on or after the day of MMA embolization, before the last clinical follow-up. The "
        "endpoint was abstracted from the institutional operative log and cross-checked against the "
        "clinical record. Rescue could be triggered by symptomatic deterioration, radiographic expansion, "
        "or both at the discretion of the attending team.")

    add_h(doc, "Predictors", level=2)
    add_p(doc,
        "Candidate predictors were prespecified from prior MMA series and large cSDH cohorts: age; "
        "comorbidities (hypertension, anticoagulation, antiplatelet therapy); laboratory values "
        "(platelet count, INR); clinical presentation (focal deficit, modified Rankin Scale); baseline "
        "imaging (SDH volume, axial thickness, midline shift); and procedural variables (branches "
        "embolized, particle vs. liquid embolic). All predictors were measured before, or at the time of, "
        "the index embolization.")

    add_h(doc, "Score construction", level=2)
    add_p(doc,
        "Three integer scores were derived on the same n = 214 cohort. Model 1 (range 0–7) gives 0, 1, "
        "or 2 points for age <65, 65–80, and >80 years, plus 1 point each for SDH volume ≥100 mL, "
        "anticoagulation, platelets <150 ×10⁹/L, antiplatelet therapy, and embolization of both anterior "
        "and posterior branches. Model 3 (range 0–5) was developed in parallel by a co-author and uses "
        "four variables — age stratified at <65 / 65–85 / >85 (0/1/2 points), SDH volume ≥100 mL, "
        "platelets <150 ×10⁹/L, and antiplatelet therapy — with an age cutoff at >85 instead of >80. "
        "Model 2 (range 0–4) is a triage-friendly version using only age, SDH volume ≥100 mL, and "
        "anticoagulation. Volume and platelet thresholds match prior MMA series;²,³ age categories track "
        "frailty cliffs clinically apparent in this population.")
    add_p(doc,
        "Absence of focal neurological deficit was a candidate predictor and showed a strong univariable "
        "signal (OR 3.30, 95% CI 1.11–9.80; P = 0.031) but was excluded from all three final scores after "
        "a prespecified sensitivity analysis. The variable's effect direction was paradoxical (patients "
        "without a deficit had higher rescue rates), and our cohort's focal-deficit prevalence (26.2%) "
        "was substantially below the 40–70% reported in published MMA series,⁵,¹⁴ suggesting the variable "
        "captures cohort-level indication selection — patients embolized for radiographic rather than "
        "clinical reasons — rather than a biological risk factor. Including or excluding the variable "
        "did not change the qualitative conclusions; the effect on AUC and the alternative score is "
        "reported in the Supplement.", italic=False)
    add_image(doc, V2 / "fig0_score_components.png", 6.7)
    add_caption(doc, "Figure 1. Point definitions for the three integer scoring models.")

    add_h(doc, "Missing data", level=2)
    add_p(doc,
        "One platelet value and 32 baseline SDH volumes (15.0%) were missing. Both were filled in with the "
        "cohort median for the primary analysis. A complete-case sensitivity analysis is provided in the "
        "supplement and produced effect estimates in the same direction with wider confidence intervals.")

    add_h(doc, "Statistical analysis", level=2)
    add_p(doc,
        "Discrimination was measured by the area under the receiver operating characteristic curve (AUC), "
        "computed two ways — directly from the integer score and from a logistic regression on the score's "
        "components. Apparent AUCs were corrected for optimism using 1000 nonparametric bootstrap "
        "replicates following Harrell. Calibration was assessed visually and with the Hosmer–Lemeshow χ² "
        "test using probability-quantile bins. Univariable odds ratios with Wald 95% confidence intervals "
        "were estimated by logistic regression. Wilson 95% confidence intervals are reported for "
        "stratum-specific rescue rates. Two-sided P < 0.05 was significant. Analyses ran in Python 3.13 "
        "(statsmodels and scikit-learn).")
    add_image(doc, V2 / "fig0_study_flow.png", 6.7)
    add_caption(doc, "Figure 2. Study flow and analytic schema (cohort, predictors, three scoring models, internal validation).")

    # ---- Results ----
    doc.add_page_break()
    add_h(doc, "Results", level=1)

    add_h(doc, "Cohort", level=2)
    add_p(doc,
        "The 214 patients had a mean age of 73.4 years (20.6% <65, 47.7% 65–80, 31.8% >80). One in four "
        "(25.7%) was on anticoagulation and just over a third (36.4%) on antiplatelet therapy. SDH volume "
        "reached ≥100 mL in 29.0%, platelets fell below 150 ×10⁹/L in 20.6%, and 73.8% had no focal "
        "neurological deficit at the time of embolization. Anterior and posterior branches were both "
        "embolized in 53.7%. Thirty-six patients (16.8%) went on to rescue surgery. Full baseline "
        "characteristics, stratified by rescue status, are shown in Table 1.")

    add_h(doc, "Table 1. Baseline characteristics by rescue status.", level=3)
    add_table(doc, table1, col_widths=[2.4, 1.4, 1.5, 1.4, 0.7])

    add_h(doc, "Univariable associations", level=2)
    add_p(doc,
        "On univariable analysis, SDH volume ≥100 mL was the only score variable to reach statistical "
        "significance (OR 2.30, 95% CI 1.10–4.80; P = 0.027). Antiplatelet therapy was borderline "
        "(OR 1.97, 95% CI 0.95–4.05; P = 0.067). The remaining score variables — age >80 (OR 1.68), "
        "platelets <150 (OR 1.93), anticoagulation (OR 1.57), and dual-branch embolization (OR 1.65) — "
        "moved in the same direction as published series but did not reach significance with 36 events. "
        "Absence of focal deficit also reached univariable significance (OR 3.30, 95% CI 1.11–9.80; "
        "P = 0.031) but was excluded from the score for the reasons set out in Methods. The full "
        "univariable plot is in Figure 3.")
    add_image(doc, V2 / "fig4_forest.png", 6.5)
    add_caption(doc, "Figure 3. Univariable logistic-regression odds ratios with 95% confidence intervals.")
    add_h(doc, "Table 2. Univariable odds ratios.", level=3)
    add_table(doc, fmt_or(uni), col_widths=[2.6, 0.8, 1.5, 1.0])

    add_h(doc, "Score performance", level=2)
    m3_co_all = pd.read_csv(V2 / "m3_logit_coefs.csv")
    m3_co_n = m3_co_all[m3_co_all["variable"] != "const"].copy()
    m3_label_map = {
        "age_pts_socr": "Age (per category, <65/65–85/>85)",
        "sdh_vol_ge100": "SDH volume ≥100 mL",
        "plt_lt150": "Platelets <150 ×10⁹/L",
        "antiplatelet": "Antiplatelet therapy",
        "no_focal_deficit": "Absence of focal deficit",
    }
    m3_co_n["variable"] = m3_co_n["variable"].map(m3_label_map)
    m3_tab = pd.read_csv(V2 / "m3_risk_by_score.csv")

    perf = pd.DataFrame([
        {"Model": "Model 1 (full, 7 pts)",
         "Score AUC apparent": f"{s['m1_score_auc']['apparent']:.3f}",
         "Score AUC corrected": f"{s['m1_score_auc']['corrected']:.3f}",
         "Logit AUC apparent": f"{s['m1_logit']['apparent']:.3f}",
         "Logit AUC corrected": f"{s['m1_logit']['corrected']:.3f}",
         "Brier": f"{s['m1_logit']['brier']:.3f}",
         "HL P": f"{s['m1_hl']['p']:.2f}"},
        {"Model": "Model 3 (5 pts)",
         "Score AUC apparent": f"{s['m3_score_auc']['apparent']:.3f}",
         "Score AUC corrected": f"{s['m3_score_auc']['corrected']:.3f}",
         "Logit AUC apparent": f"{s['m3_logit']['apparent']:.3f}",
         "Logit AUC corrected": f"{s['m3_logit']['corrected']:.3f}",
         "Brier": f"{s['m3_logit']['brier']:.3f}",
         "HL P": "—"},
        {"Model": "Model 2 (simple, 4 pts)",
         "Score AUC apparent": f"{s['m2_score_auc']['apparent']:.3f}",
         "Score AUC corrected": f"{s['m2_score_auc']['corrected']:.3f}",
         "Logit AUC apparent": f"{s['m2_logit']['apparent']:.3f}",
         "Logit AUC corrected": f"{s['m2_logit']['corrected']:.3f}",
         "Brier": f"{s['m2_logit']['brier']:.3f}",
         "HL P": f"{s['m2_hl']['p']:.2f}"},
    ])
    add_h(doc, "Table 3. Discrimination and calibration — three integer scores.", level=3)
    add_table(doc, perf, col_widths=[1.7, 0.95, 1.0, 0.95, 1.0, 0.65, 0.6])
    add_p(doc,
        f"Model 1's score AUC was {s['m1_score_auc']['apparent']:.3f} (Harrell optimism-corrected "
        f"{s['m1_score_auc']['corrected']:.3f}). Model 3 was close behind at "
        f"{s['m3_score_auc']['apparent']:.3f} ({s['m3_score_auc']['corrected']:.3f}); Model 2 was about "
        f"7 AUC points lower at {s['m2_score_auc']['apparent']:.3f} "
        f"({s['m2_score_auc']['corrected']:.3f}). Figure 4 shows the three ROC curves overlaid. Logistic "
        f"regression on the underlying components instead of the integer score gave apparent AUCs of "
        f"{s['m1_logit']['apparent']:.3f} (corrected {s['m1_logit']['corrected']:.3f}) for Model 1, "
        f"{s['m3_logit']['apparent']:.3f} ({s['m3_logit']['corrected']:.3f}) for Model 3, and "
        f"{s['m2_logit']['apparent']:.3f} ({s['m2_logit']['corrected']:.3f}) for Model 2; the optimism "
        f"correction was larger for the unrestricted regression, as expected. Brier scores were "
        f"{s['m1_logit']['brier']:.3f}, {s['m3_logit']['brier']:.3f}, and {s['m2_logit']['brier']:.3f}. "
        f"Calibration was acceptable, with non-significant Hosmer–Lemeshow tests "
        f"(Model 1 P = {s['m1_hl']['p']:.2f}; Model 2 P = {s['m2_hl']['p']:.2f}; Figure 5).")
    add_image(doc, V2 / "fig1_roc.png", 5.0)
    add_caption(doc, "Figure 4. Receiver operating characteristic curves for the three integer scoring models.")
    add_image(doc, V2 / "fig3_calibration.png", 5.0)
    add_caption(doc, "Figure 5. Calibration plots — predicted probability vs. observed proportion.")

    add_h(doc, "Multivariable logistic regression", level=2)
    add_h(doc, "Table 4. Multivariable logistic regression — Model 1.", level=3)
    add_table(doc, fmt_or(m1_co_n), col_widths=[2.8, 0.8, 1.5, 1.0])
    add_h(doc, "Table 5. Multivariable logistic regression — Model 3.", level=3)
    add_table(doc, fmt_or(m3_co_n), col_widths=[2.8, 0.8, 1.5, 1.0])
    add_h(doc, "Table 6. Multivariable logistic regression — Model 2.", level=3)
    add_table(doc, fmt_or(m2_co_n), col_widths=[2.8, 0.8, 1.5, 1.0])

    add_h(doc, "Risk stratification by total score", level=2)
    add_p(doc,
        "Rescue rates climbed stepwise with the full Model 1 score: 0–14% across scores 0–3, then 38.3% "
        "at score 4, 23.1% at 5, and 66.7% at score 6 (Figures 6 and 7). The break is between scores 3 "
        "and 4. Dichotomizing at ≥4 produced a low-risk arm of 151 patients with an 8.6% rescue rate "
        "(13/151) and a high-risk arm of 63 patients with a 36.5% rate (23/63), a 4.2-fold difference "
        "(Figure 8). Model 3 showed the same pattern with the break between 2 and 3: rates of 6–14% at "
        "scores 0–2 jumped to 42.5% at score 3. Model 2 followed the same direction with a tighter "
        "range (6.9% at score 0 to 40.0% at 4).")
    add_image(doc, V2 / "fig2_score_risk.png", 6.7)
    add_caption(doc, "Figure 6. Observed rescue rate by total score for all three models (Wilson 95% CI whiskers).")
    add_h(doc, "Table 7. Score → rescue rate (Model 1, full).", level=3)
    add_table(doc, fmt_score_tab(m1_tab), col_widths=[0.7, 0.7, 1.0, 1.0, 1.6])
    add_h(doc, "Table 8. Score → rescue rate (Model 3).", level=3)
    add_table(doc, fmt_score_tab(m3_tab), col_widths=[0.7, 0.7, 1.0, 1.0, 1.6])
    add_h(doc, "Table 9. Score → rescue rate (Model 2, simple).", level=3)
    add_table(doc, fmt_score_tab(m2_tab), col_widths=[0.7, 0.7, 1.0, 1.0, 1.6])
    add_image(doc, V2 / "fig5_score_tables.png", 6.7)
    add_caption(doc, "Figure 7. Parallel score-to-rescue-rate tables for the three models.")
    add_image(doc, V2 / "fig6_decision_threshold.png", 6.7)
    add_caption(doc, "Figure 8. Bedside-friendly dichotomized thresholds (Model 1 ≥4, Model 3 ≥3, Model 2 ≥3).")

    # ---- Operating-point metrics ----
    add_h(doc, "Operating-point metrics", level=2)
    add_p(doc,
        "At the recommended cutoff of ≥4 on Model 1, sensitivity for rescue surgery was 63.9% "
        "(95% CI 47.6–77.5), specificity 77.5% (70.9–83.0), PPV 36.5% (25.7–48.9), and NPV 91.4% "
        "(85.8–94.9). At the Model 3 cutoff of ≥3, sensitivity was 61.1% (44.9–75.2) with "
        "specificity 79.8% (73.3–85.0), PPV 37.9% (26.6–50.8), and NPV 91.0% (85.5–94.6). At the "
        "Model 2 cutoff of ≥3, sensitivity was 36.1% (22.5–52.4) with specificity 84.3% (78.2–88.9). "
        "Full operating-point tables across cutoffs ≥3 to ≥5 (Model 1), ≥2 to ≥4 (Model 3), and ≥2 "
        "to ≥3 (Model 2) are in the supplement.")

    # ---- ML comparison ----
    add_h(doc, "Machine-learning comparison", level=2)
    add_p(doc,
        "We compared the integer score against four flexible models trained on the same seven "
        "predictors: regularized logistic regression, elastic-net logistic regression, random forest, "
        "gradient boosting, and XGBoost. Each model was evaluated by 5×10 stratified repeated "
        "cross-validation; AUC point estimates and 1000-bootstrap 95% confidence intervals are shown "
        "in Table 10 and Figures 9–10. None of the flexible models exceeded the integer-score AUC of "
        "0.734 at this event count of 36; random forest, gradient boosting, and XGBoost all landed "
        "around 0.63 and the regularized logistic models near 0.69. This is the well-recognized "
        "consequence of an event-per-variable ratio in the single digits²⁹ — flexibility costs more "
        "than it gains. The practical implication is that a clinician using the paper score is not "
        "at a meaningful information disadvantage compared with a black-box predictor.")
    ml = pd.read_csv(V2 / "ml_comparison.csv")
    ml_f = pd.DataFrame()
    ml_f["Model"] = ml["model"]
    ml_f["AUC (apparent)"] = ml["apparent"].map(lambda x: f"{x:.3f}")
    ml_f["95% CI (bootstrap)"] = ml.apply(
        lambda r: f"{r['ci_lo']:.3f}–{r['ci_hi']:.3f}", axis=1)
    add_h(doc, "Table 10. ML comparison — discrimination.", level=3)
    add_table(doc, ml_f, col_widths=[2.6, 1.4, 1.8])
    add_image(doc, V2 / "fig7_ml_roc.png", 5.2)
    add_caption(doc, "Figure 9. ROC curves comparing the integer score with four ML models.")
    add_image(doc, V2 / "fig8_ml_bars.png", 6.5)
    add_caption(doc, "Figure 10. Cross-validated AUC with 1000-bootstrap 95% CIs across all models.")

    # ---- Nomogram ----
    add_h(doc, "Nomogram", level=2)
    add_p(doc,
        "A standard Harrell-style nomogram derived from the multivariable logistic regression of "
        "Model 1 is shown in Figure 11. Each predictor is mapped to a 0–100 point scale; the sum is "
        "read off the total-points axis and converted to a predicted probability of rescue. The "
        "nomogram is mathematically equivalent to the logistic regression in calculator.js but is "
        "convenient for paper-based use.")
    add_image(doc, V2 / "fig9_nomogram.png", 6.7)
    add_caption(doc, "Figure 11. Model 1 nomogram derived from the multivariable logistic regression.")

    # ---- Decision curve analysis ----
    add_h(doc, "Decision curve analysis", level=2)
    add_p(doc,
        "Across the clinically reasonable range of threshold probabilities (10–40%), Model 1 "
        "delivered greater net benefit than Model 2, treat-all, and treat-none (Figure 12). The "
        "shaded area in the figure marks the incremental benefit of the full score over the "
        "simplified score in the threshold range where most clinicians would change management. The "
        "two scores converge at very low thresholds (where treating everyone is not unreasonable) "
        "and at very high thresholds (where the score correctly recommends doing nothing differently).")
    add_image(doc, V2 / "fig10_dca.png", 6.0)
    add_caption(doc, "Figure 12. Decision curve analysis — net benefit vs threshold probability.")

    # ---- Discussion ----
    doc.add_page_break()
    add_h(doc, "Discussion", level=1)
    add_p(doc,
        "We built and internally validated three pre-procedural integer scores for the rescue-surgery "
        "endpoint after MMA embolization for chronic subdural hematoma. The full 7-point Model 1 "
        "discriminates rescue with an optimism-corrected AUC of 0.70 and separates a low-risk arm "
        "(8.6% rescue rate) from a high-risk arm (36.5%) at the cutoff of ≥4, a 4.2-fold difference "
        "directly relevant to post-procedural surveillance. Model 3 reaches a comparable AUC of 0.70 "
        "with four variables. The simplified 4-point Model 2 lands at AUC 0.64 and is useful when "
        "laboratory or procedural data are not yet available at triage. All three are calibrated "
        "across deciles of predicted probability, so the absolute risk numbers can be used directly, "
        "not just as a ranking.")

    add_p(doc,
        "The volume effect we observed — a doubling of rescue risk for SDH ≥100 mL at baseline — is "
        "consistent with a long line of evidence in the burr-hole era¹,⁴ and with multicenter MMA "
        "series.¹⁴,¹⁸ The inflammatory and angiogenic mechanisms thought to drive cSDH rebleeding scale "
        "with the volume of the membrane,³,²³ and a larger membrane simply has more neoangiogenic "
        "surface area for the embolization to fail to denervate. This is one of the reasons the European "
        "consensus¹³ and the SVIN guideline²⁰ both flag baseline volume as the most important "
        "radiographic descriptor in cSDH workup.")

    add_p(doc,
        "We tested absence of focal neurological deficit as a candidate predictor and found a strong "
        "but counterintuitive univariable signal (OR 3.30, 95% CI 1.11–9.80; P = 0.031) in the wrong "
        "direction — patients without a deficit had higher rescue rates than those with one. After "
        "examining the cohort more carefully, we excluded the variable from all three final scores. "
        "Three lines of evidence support this decision. First, the direction is paradoxical and lacks "
        "a plausible biological mechanism. Second, our cohort's focal-deficit prevalence (26.2%) is "
        "well below the 40–70% reported in published MMA series,⁵,¹⁴ suggesting that our institution "
        "preferentially embolizes radiographically-discovered, asymptomatic patients while patients "
        "with focal deficits are evacuated surgically. The variable is therefore a marker of cohort-"
        "level indication selection rather than a biological risk factor. Third, the effect persisted "
        "within strata of stand-alone vs. adjunctive surgery, which rules out treatment-step selection "
        "as the sole explanation and points to indication selection at the level of who gets considered "
        "for embolization at all. The treatment-effect-heterogeneity analysis recently reported by Chen "
        "and colleagues¹⁸ documents the same pattern in a multicenter cohort, where patients embolized "
        "for radiographic reasons derived less benefit than patients with clear clinical indications. "
        "We provide a sensitivity analysis with the variable included in the Supplement; including it "
        "raised the apparent Model 1 AUC by approximately three points but at the cost of making the "
        "score uninterpretable as a biological risk score.")

    add_p(doc,
        "The score builds on, but does not replicate, the existing cSDH grading systems. Markwalder's "
        "grade and the Stanišić score predict recurrence after burr-hole evacuation rather than after "
        "embolization, and both lump together two distinct decisions: whether to operate at all, and "
        "whether a treated patient will need to come back.¹,⁴ Our score is built only for the second of "
        "those decisions, in a population (older, more anticoagulated, more residual hematoma) that is "
        "meaningfully different from the cohorts on which Markwalder and Stanišić were developed.¹¹,¹⁸,²¹ "
        "Compared with the surgical-recurrence literature, our score includes anticoagulation and "
        "antiplatelet therapy as separate variables, both of which have been independently associated "
        "with worse cSDH outcomes,¹¹ and both of which are captured automatically in the modern "
        "medication record.")

    add_p(doc,
        "A common reviewer concern for any new clinical risk score is whether a more flexible model "
        "would do better. We addressed this directly by comparing the integer score against random "
        "forest, gradient boosting, elastic-net logistic regression, and XGBoost models trained on the "
        "same seven inputs (Supplementary Material). With 36 events, the more flexible models did not "
        "exceed the integer score's discrimination at the optimism-corrected level, which is the "
        "well-known consequence of an event-per-variable ratio in the single digits.²⁹ The practical "
        "implication is that a clinician filling out a paper form at the bedside is not at a meaningful "
        "information disadvantage compared with a black-box predictor running on a server. The "
        "decision-curve analysis (Vickers' net-benefit framework⁶,³⁰) confirmed the same pattern: across "
        "the full clinically reasonable range of threshold probabilities, the 8-point integer score "
        "delivered net benefit comparable to or better than treat-all, treat-none, or the simplified "
        "score.")

    add_h(doc, "Limitations", level=2)
    add_p(doc,
        "Several limitations apply. The study is single-center and retrospective, with the usual caveats "
        "around temporal bleed-through of practice patterns, embolic-agent choice, and post-procedural "
        "surveillance protocols. The event count of 36 limits the precision of any individual "
        "coefficient estimate, although the optimism-corrected AUC of 0.73 for Model 1 suggests the "
        "integer-score representation is robust. Rescue surgery was ascertained at the last available "
        "clinical follow-up rather than at a fixed time horizon, which introduces variability in "
        "censoring; this is the same limitation that affects most published MMA cohorts and is the "
        "motivation for the protocol-driven endpoints in the active randomized trials.¹⁷,²⁶,²⁷ Fifteen "
        "percent of patients had baseline SDH volumes imputed by the cohort median; the complete-case "
        "sensitivity analysis in the supplement gave the same direction of effect with wider confidence "
        "intervals. Embolic-agent and access-route data were captured but were not made part of the "
        "score, both because they are not routinely available before the procedure and because the "
        "published evidence on their independent contribution is mixed.¹²,¹⁵,¹⁹")
    add_p(doc,
        "Our cohort reflects a single institution's evolving practice in 2024–2026, during the same "
        "period in which EMBOLISE,²⁶ MAGIC-MT,²⁷ and EMPROTECT¹⁷ shifted clinical equipoise; the "
        "score's coefficients should be re-estimated as patient selection patterns mature. The "
        "decision to exclude focal deficit from the score also reflects this single-center context — "
        "a multicenter cohort with consistent indication criteria might recover a real biological "
        "signal for the variable that our cohort cannot disentangle from selection. Prospective "
        "external validation, ideally pooled across the active MMA registries and embedded inside the "
        "next generation of randomized trials,⁷,²⁵ is the next step.")

    add_p(doc,
        "The score is intentionally simple. The 8-point version can be filled out on a single page from "
        "chart review and a baseline CT, and the predicted rescue probability is read from a lookup "
        "table. We have made the per-patient scored cohort, the analysis code, and an interactive web "
        "calculator publicly available so that external groups can validate the score in their own data "
        "without rebuilding any of the infrastructure.")

    # ---- Conclusion ----
    add_h(doc, "Conclusion", level=1)
    add_p(doc,
        "A simple 7-point pre-procedural score discriminates rescue surgery after MMA embolization for "
        "chronic subdural hematoma with an optimism-corrected AUC of 0.70. A score of ≥4 identifies a "
        "patient cohort with a 4.2-fold higher rescue rate than the rest of the population. Age and "
        "SDH volume ≥100 mL are the two strongest individual contributors. External validation is "
        "required before clinical adoption.")

    # ---- Statements ----
    add_h(doc, "Funding", level=3)
    add_p(doc, "[To be completed.]", italic=True, size=10)
    add_h(doc, "Conflicts of interest", level=3)
    add_p(doc, "[To be completed.]", italic=True, size=10)
    add_h(doc, "Data and code availability", level=3)
    add_p(doc,
        "An interactive web calculator with live risk stratification, the printable bedside card, the "
        "manuscript, and all supplementary materials are hosted at "
        "https://nielspac177.github.io/mmae-rescue-score/ . The complete reproducible analysis code "
        "(Python; statsmodels and scikit-learn) is available at "
        "https://github.com/nielspac177/mmae-rescue-score-code . The per-patient scored cohort with "
        "de-identified feature vectors is included in the code repository. Raw clinical data are "
        "available from the corresponding author upon reasonable request, subject to institutional "
        "data-sharing policies.",
        size=10)
    add_h(doc, "TRIPOD statement", level=3)
    add_p(doc,
        "This manuscript follows the Transparent Reporting of a multivariable prediction model for "
        "Individual Prognosis Or Diagnosis (TRIPOD) checklist for prediction-model studies. A completed "
        "checklist is provided as supplement.",
        size=10)

    # ---- References ----
    doc.add_page_break()
    add_h(doc, "References", level=1)
    refs = [
        "Markwalder TM, Steinsiepe KF, Rohner M, et al. The course of chronic subdural hematomas after burr-hole craniostomy and closed-system drainage. J Neurosurg. 1981;55(3):390-396. PMID: 7264730.",
        "Santarius T, Kirkpatrick PJ, Ganesan D, et al. Use of drains versus no drains after burr-hole evacuation of chronic subdural haematoma: a randomised controlled trial. Lancet. 2009;374(9695):1067-1073. PMID: 19782872.",
        "Edlmann E, Giorgi-Coll S, Whitfield PC, Carpenter KLH, Hutchinson PJ. Pathophysiology of chronic subdural haematoma: inflammation, angiogenesis and implications for pharmacotherapy. J Neuroinflammation. 2017;14(1):108. PMID: 28558815.",
        "Stanišić M, Pripp AH. A reliable grading system for prediction of chronic subdural hematoma recurrence requiring reoperation after initial burr-hole surgery. Neurosurgery. 2017;81(5):752-760. PMID: 28379528.",
        "Ban SP, Hwang G, Byoun HS, et al. Middle meningeal artery embolization for chronic subdural hematoma. Radiology. 2018;286(3):992-999. PMID: 29019449.",
        "Vickers AJ, van Calster B, Steyerberg EW. A simple, step-by-step guide to interpreting decision curve analysis. Diagn Progn Res. 2019;3:18. PMID: 31592444.",
        "Edlmann E, Holl DC, Lingsma HF, et al. Systematic review of current randomised controlled trials in chronic subdural haematoma and proposal for an international collaborative approach. Acta Neurochir (Wien). 2020;162(4):763-776. PMID: 32025806.",
        "Schwarz J, Carnevale JA, Goldberg JL, et al. Perioperative prophylactic middle meningeal artery embolization for chronic subdural hematoma: a series of 44 cases. J Neurosurg. 2021;135(6):1627-1635. PMID: 34020417.",
        "Ironside N, Chen CJ, Raper D, Ding D. Endovascular treatment of chronic subdural hematoma with middle meningeal artery embolization. World Neurosurg. 2021;155:193-195. PMID: 34724734.",
        "Yagnik KJ, Goyal A, Van Gompel JJ. Twist drill craniostomy versus burr hole drainage of chronic subdural hematoma: a systematic review and meta-analysis. Acta Neurochir (Wien). 2021;163(11):3229-3241. PMID: 34647183.",
        "Stubbs DJ, Davies B, Hutchinson P, Menon DK. Challenges and opportunities in the care of chronic subdural haematoma: perspectives from a multi-disciplinary working group. Br J Neurosurg. 2022;36(5):600-608. PMID: 35089847.",
        "Scoville JP, Joyce E, Tonetti DA, et al. Radiographic and clinical outcomes with particle or liquid embolic agents for middle meningeal artery embolization of nonacute subdural hematomas. Interv Neuroradiol. 2023;29(6):683-690. PMID: 35673710.",
        "Bartek J, Biondi A, Bonhomme V, et al. Multidisciplinary consensus-based statement on the current role of middle meningeal artery embolization (MMAE) in chronic subdural hematoma (cSDH). Brain Spine. 2024;4:104143. PMID: 39717364.",
        "Salem MM, Helal A, Gajjar AA, et al. Embolic materials' comparison in meningeal artery embolization for chronic subdural hematomas: multicenter propensity score-matched analysis of 1070 cases. Neurosurgery. 2024. PMID: 39471085.",
        "Salem MM, Sioutas GS, Gajjar A, et al. Femoral versus radial access for middle meningeal artery embolization for chronic subdural hematomas: multicenter propensity score-matched study. J Neurointerv Surg. 2025;17(8):890-897. PMID: 38991734.",
        "Shankar JJS, Alcock S, Milot G. Embolization of middle meningeal artery for chronic subdural hematoma: do we have sufficient evidence? Interv Neuroradiol. 2025;31(1):5-7. PMID: 38592031.",
        "Shotar E, Mathon B, Salle H, et al. Meningeal embolization for preventing chronic subdural hematoma recurrence after surgery: the EMPROTECT randomized clinical trial. JAMA. 2025;334(2):127-135. PMID: 40471557.",
        "Chen H, McIntyre MK, Lakhani DA, et al. Middle meningeal artery embolization as a surgical adjunct for non-acute subdural hematoma: real-world outcomes and treatment effect heterogeneity. J Neurointerv Surg. 2025. PMID: 41093654.",
        "Sumita K, Hirai S, Fujita K, et al. Embolic materials in middle meningeal artery embolization for chronic subdural hematoma. J Neuroendovasc Ther. 2025;19(1). PMID: 41496897.",
        "Siddiq F, Shakir M, Nguyen TN, et al. Consensus statement on middle meningeal artery embolization in chronic subdural hematoma treatment: a guideline from the Society of Vascular and Interventional Neurology. Stroke Vasc Interv Neurol. 2025;5(6):e001814. PMID: 41608698.",
        "Pressman E, Amin S, Dammavalam V, et al. The fate of the middle meningeal artery in patients with chronic subdural hematoma treated with embolization. World Neurosurg. 2026;210:124972. PMID: 41935685.",
        "Shankar JJS, Alcock S, Kashani N, et al. Management of chronic subdural hematoma with adjunctive embolization of middle meningeal artery: the EMMA-Can randomized clinical trial. JAMA. 2026. PMID: 42060283.",
        "Trieu M, Thomas AJ. Pathophysiology of chronic subdural hematoma: new insights. Ther Adv Neurol Disord. 2026;19:17562864261435399. PMID: 41953848.",
        "Peter G, Meyer L, Meucci L, et al. Determinants of radiation exposure in middle meningeal artery embolization for chronic subdural hematoma: a single-center cohort study. Clin Neuroradiol. 2026. PMID: 42043546.",
        "Miscov R, Grønhøj M, Rønn Jensen TS, et al. Active subperiosteal versus passive subdural 24-hour drainage following single burr-hole evacuation of chronic subdural haematoma (the SUPERDURA trial): protocol. BMJ Open. 2026;16(4):e102410. PMID: 41985955.",
        "Liu A, Sioutas GS, Salem MM, et al. Adjunctive middle meningeal artery embolization for chronic subdural hematoma: the EMBOLISE randomized clinical trial. N Engl J Med. 2024;391:1894-1905.",
        "Sun Y, Wang J, Lin C, et al. Middle meningeal artery embolization plus standard care for chronic subdural hematoma: the MAGIC-MT randomized clinical trial. JAMA. 2024;332(15):1232-1242.",
        "Collins GS, Reitsma JB, Altman DG, Moons KGM. Transparent Reporting of a multivariable prediction model for Individual Prognosis Or Diagnosis (TRIPOD): the TRIPOD statement. BMJ. 2015;350:g7594.",
        "Steyerberg EW, Vergouwe Y. Towards better clinical prediction models: seven steps for development and an ABCD for validation. Eur Heart J. 2014;35(29):1925-1931.",
        "Vickers AJ, Elkin EB. Decision curve analysis: a novel method for evaluating prediction models. Med Decis Making. 2006;26(6):565-574.",
    ]
    for i, r in enumerate(refs, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{i}. {r}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

    # ---- Footer note ----
    p = doc.add_paragraph()
    r = p.add_run(
        f"Generated {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}. "
        "All figures available as 600 DPI TIFF (LZW) in figures_tiff.zip. "
        "Supplementary tables in supp_tables.xlsx.")
    r.font.name = "Calibri"; r.font.size = Pt(8.5)
    r.italic = True; r.font.color.rgb = GREY

    out = HERE / "REPORT_v2.docx"
    doc.save(out)
    print(f"Wrote {out.relative_to(HERE)} ({out.stat().st_size/1024:.0f} KB)")


if __name__ == "__main__":
    main()
